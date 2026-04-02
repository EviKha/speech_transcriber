from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from faster_whisper import WhisperModel


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Transcribe speech from an audio/video file into TXT or DOCX."
    )
    parser.add_argument("input", help="Path to the input audio/video file.")
    parser.add_argument(
        "-o",
        "--output",
        help="Path to the output file. Defaults to the input name with .txt or .docx extension.",
    )
    parser.add_argument(
        "--format",
        choices=("txt", "docx"),
        default="txt",
        help="Output file format. Default: txt.",
    )
    parser.add_argument(
        "--model",
        default="small",
        help="Whisper model size: tiny, base, small, medium, large-v3. Default: small.",
    )
    parser.add_argument(
        "--language",
        default=None,
        help="Optional language code, for example ru or en. If omitted, language is auto-detected.",
    )
    parser.add_argument(
        "--device",
        choices=("auto", "cpu", "cuda"),
        default="auto",
        help="Inference device. Default: auto.",
    )
    parser.add_argument(
        "--compute-type",
        default=None,
        help=(
            "Optional compute type override, for example int8, int8_float32, "
            "float16 or float32."
        ),
    )
    parser.add_argument(
        "--with-timestamps",
        action="store_true",
        help="Include segment timestamps in the output text.",
    )
    return parser.parse_args()


def build_output_path(input_path: Path, output: str | None, output_format: str) -> Path:
    if output:
        return Path(output).expanduser().resolve()
    return input_path.with_suffix(f".{output_format}")


def format_timestamp(seconds: float) -> str:
    total_ms = int(seconds * 1000)
    hours, remainder = divmod(total_ms, 3_600_000)
    minutes, remainder = divmod(remainder, 60_000)
    secs, ms = divmod(remainder, 1000)
    return f"{hours:02}:{minutes:02}:{secs:02}.{ms:03}"


def collect_text(segments, with_timestamps: bool) -> str:
    lines: list[str] = []
    for segment in segments:
        text = segment.text.strip()
        if not text:
            continue
        if with_timestamps:
            start = format_timestamp(segment.start)
            end = format_timestamp(segment.end)
            lines.append(f"[{start} - {end}] {text}")
        else:
            lines.append(text)
    return "\n".join(lines).strip() + "\n"


def save_txt(output_path: Path, text: str) -> None:
    output_path.write_text(text, encoding="utf-8")


def save_docx(output_path: Path, text: str, title: str) -> None:
    document = Document()
    document.add_heading("Speech Transcription", level=1)
    document.add_paragraph(f"Source file: {title}")
    for block in text.strip().splitlines():
        document.add_paragraph(block)
    document.save(output_path)


def resolve_compute_type(device: str, compute_type: str | None) -> str:
    if compute_type:
        return compute_type
    if device == "cuda":
        return "float16"
    return "int8"


def main() -> int:
    args = parse_args()

    input_path = Path(args.input).expanduser().resolve()
    if not input_path.exists():
        print(f"Input file not found: {input_path}", file=sys.stderr)
        return 1

    output_path = build_output_path(input_path, args.output, args.format)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    device = "cpu" if args.device == "auto" else args.device
    compute_type = resolve_compute_type(device, args.compute_type)

    try:
        model = WhisperModel(args.model, device=device, compute_type=compute_type)
        segments, info = model.transcribe(
            str(input_path),
            language=args.language,
            vad_filter=True,
        )
        text = collect_text(segments, args.with_timestamps)
    except Exception as exc:
        print("Transcription failed.", file=sys.stderr)
        print(str(exc), file=sys.stderr)
        return 2

    if args.format == "txt":
        save_txt(output_path, text)
    else:
        save_docx(output_path, text, input_path.name)

    detected_language = getattr(info, "language", "unknown")
    duration = getattr(info, "duration", None)
    print(f"Saved: {output_path}")
    print(f"Detected language: {detected_language}")
    if duration is not None:
        print(f"Duration: {duration:.1f} sec")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
