from __future__ import annotations

import os
import platform
import queue
import subprocess
import tempfile
import threading
import time
import traceback
import tkinter as tk
import warnings
import wave
from dataclasses import dataclass
from pathlib import Path
from tkinter import filedialog, messagebox, ttk

import numpy as np
import soundcard as sc
from docx import Document
from faster_whisper import WhisperModel

try:
    from resemblyzer import VoiceEncoder, preprocess_wav
except ImportError:
    VoiceEncoder = None
    preprocess_wav = None

try:
    import simpleaudio as sa
except ImportError:
    sa = None


APP_TITLE = "Live Meeting Transcriber"
SAMPLE_RATE = 16_000
DEFAULT_RECORD_SECONDS = 3
STABLE_RECORD_SECONDS = 5
DEFAULT_BLOCK_FRAMES = 2048
STABLE_BLOCK_FRAMES = 4096
SEARCH_TAG = "search_match"
CURRENT_SEARCH_TAG = "search_current"
LINE_PREFIX = "line_"
MIN_DIARIZATION_SECONDS = 0.8
SPEAKER_MATCH_THRESHOLD = 0.72
PLAY_ICON = "▶"
IMPORTANT_ICON = "❗"
TASK_ICON = "✓"
QUESTION_ICON = "?"
EMPTY_MARKER_ICON = "○"
CLIP_PADDING_BEFORE = 0.12
CLIP_PADDING_AFTER = 0.18

os.environ.setdefault("HF_HUB_DISABLE_SYMLINKS_WARNING", "1")
warnings.filterwarnings(
    "ignore",
    message="data discontinuity in recording",
)
warnings.filterwarnings(
    "ignore",
    message="`huggingface_hub` cache-system uses symlinks by default",
)


def timestamp_now() -> str:
    return time.strftime("%Y-%m-%d %H:%M:%S")


def normalize_text(text: str) -> str:
    return " ".join(text.split()).strip()


def slugify_source(source_tag: str) -> str:
    return source_tag.lower().replace(" ", "_")


def clamp_audio(chunk: np.ndarray) -> np.ndarray:
    return np.clip(chunk, -1.0, 1.0)


def pcm16_bytes(chunk: np.ndarray) -> bytes:
    pcm = (clamp_audio(chunk) * 32767.0).astype(np.int16)
    return pcm.tobytes()


def write_wav_file(path: Path, chunk: np.ndarray) -> None:
    with wave.open(str(path), "wb") as wav_handle:
        wav_handle.setnchannels(1)
        wav_handle.setsampwidth(2)
        wav_handle.setframerate(SAMPLE_RATE)
        wav_handle.writeframes(pcm16_bytes(chunk))


def marker_icon(marker: str | None) -> str:
    if marker == "IMPORTANT":
        return IMPORTANT_ICON
    if marker == "TASK":
        return TASK_ICON
    if marker == "QUESTION":
        return QUESTION_ICON
    return ""


def action_prefix() -> str:
    return f"{PLAY_ICON} {TASK_ICON} {QUESTION_ICON} {IMPORTANT_ICON}"


@dataclass
class TranscriptEntry:
    entry_id: int
    source_tag: str
    transcript_text: str
    created_at: str
    audio_path: Path | None
    audio_start: float
    audio_end: float
    marker: str | None = None
    line_number: int | None = None

    def display_text(self, show_timestamps: bool) -> str:
        marker = marker_icon(self.marker) or EMPTY_MARKER_ICON
        prefix = f"{PLAY_ICON} {TASK_ICON} {QUESTION_ICON} {IMPORTANT_ICON} {marker}"
        if show_timestamps:
            return f"{prefix} [{self.created_at}] [{self.source_tag}] {self.transcript_text}"
        return f"{prefix} [{self.source_tag}] {self.transcript_text}"

    def has_audio(self) -> bool:
        return self.audio_path is not None and self.audio_path.exists() and self.audio_end > self.audio_start


class LiveTranscriberApp:
    def __init__(self, root: tk.Tk) -> None:
        self.root = root
        self.root.title(APP_TITLE)
        self.root.geometry("1180x820")
        self.root.minsize(980, 680)

        self.audio_queue: queue.Queue[tuple[str, np.ndarray, float, float, Path | None]] = queue.Queue(
            maxsize=24
        )
        self.ui_queue: queue.Queue[tuple[str, str]] = queue.Queue()

        self.stop_event = threading.Event()
        self.capture_thread: threading.Thread | None = None
        self.mic_thread: threading.Thread | None = None
        self.worker_thread: threading.Thread | None = None
        self.playback_thread: threading.Thread | None = None

        self.model: WhisperModel | None = None
        self.diarization_encoder: VoiceEncoder | None = None
        self.last_emitted_by_source: dict[str, str] = {}
        self.transcript_parts: list[str] = []
        self.transcript_entries: list[TranscriptEntry] = []
        self.line_to_entry_id: dict[int, int] = {}
        self.search_results: list[str] = []
        self.current_search_index = -1
        self.selected_entry_id: int | None = None
        self.speaker_profiles: dict[str, tuple[np.ndarray, int]] = {}
        self.next_speaker_id = 1

        self.output_dir = Path.cwd() / "live_transcripts"
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.audio_dir = Path.cwd() / "live_audio"
        self.audio_dir.mkdir(parents=True, exist_ok=True)
        self.audio_clips_dir = self.audio_dir / "clips"
        self.audio_clips_dir.mkdir(parents=True, exist_ok=True)
        self.logs_dir = Path.cwd() / "logs"
        self.logs_dir.mkdir(parents=True, exist_ok=True)
        self.log_path = self.logs_dir / "live_transcriber.log"
        self.last_error_details = ""
        self.session_stamp = time.strftime("%Y%m%d_%H%M%S")

        self.audio_write_lock = threading.Lock()
        self.audio_wavefiles: dict[str, wave.Wave_write] = {}
        self.audio_paths: dict[str, Path] = {}

        self.language_var = tk.StringVar(value="ru")
        self.model_var = tk.StringVar(value="small")
        self.status_var = tk.StringVar(value="Ready")
        self.output_var = tk.StringVar(value=str(self.default_output_path()))
        self.log_var = tk.StringVar(value=str(self.log_path.resolve()))
        self.timestamps_var = tk.BooleanVar(value=True)
        self.include_mic_var = tk.BooleanVar(value=True)
        self.stable_mode_var = tk.BooleanVar(value=True)
        self.meeting_diarization_var = tk.BooleanVar(value=False)
        self.search_var = tk.StringVar(value="")

        self.build_ui()
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)
        self.root.after(150, self.process_ui_queue)
        self.log_message("info", "Application started")

    def build_ui(self) -> None:
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(2, weight=1)

        controls = ttk.Frame(self.root, padding=12)
        controls.grid(row=0, column=0, sticky="ew")
        controls.columnconfigure(12, weight=1)

        ttk.Label(controls, text="Language").grid(row=0, column=0, sticky="w")
        language_combo = ttk.Combobox(
            controls,
            textvariable=self.language_var,
            values=("ru", "en", "uk", "de", "fr", "es", "auto"),
            state="readonly",
            width=10,
        )
        language_combo.grid(row=0, column=1, padx=(6, 14), sticky="w")

        ttk.Label(controls, text="Model").grid(row=0, column=2, sticky="w")
        model_combo = ttk.Combobox(
            controls,
            textvariable=self.model_var,
            values=("tiny", "base", "small", "medium", "large-v3"),
            state="readonly",
            width=12,
        )
        model_combo.grid(row=0, column=3, padx=(6, 14), sticky="w")

        ttk.Checkbutton(
            controls,
            text="Show timestamps",
            variable=self.timestamps_var,
            command=self.refresh_transcript_view,
        ).grid(row=0, column=4, padx=(0, 14), sticky="w")

        ttk.Checkbutton(
            controls,
            text="Include microphone",
            variable=self.include_mic_var,
        ).grid(row=0, column=5, padx=(0, 14), sticky="w")

        ttk.Checkbutton(
            controls,
            text="Stable mode",
            variable=self.stable_mode_var,
        ).grid(row=0, column=6, padx=(0, 14), sticky="w")

        ttk.Checkbutton(
            controls,
            text="Split meeting speakers",
            variable=self.meeting_diarization_var,
        ).grid(row=0, column=7, padx=(0, 14), sticky="w")

        self.start_button = ttk.Button(controls, text="Start", command=self.start_transcription)
        self.start_button.grid(row=0, column=8, padx=(0, 8), sticky="ew")

        self.stop_button = ttk.Button(
            controls,
            text="Stop",
            command=self.stop_transcription,
            state="disabled",
        )
        self.stop_button.grid(row=0, column=9, padx=(0, 8), sticky="ew")

        self.export_button = ttk.Button(controls, text="Export DOCX", command=self.export_docx)
        self.export_button.grid(row=0, column=10, padx=(0, 8), sticky="e")

        self.summary_button = ttk.Button(controls, text="Summary", command=self.show_summary)
        self.summary_button.grid(row=0, column=11, padx=(0, 8), sticky="e")

        self.play_button = ttk.Button(controls, text="Play Selected", command=self.play_selected_entry)
        self.play_button.grid(row=0, column=12, padx=(0, 8), sticky="e")

        self.diagnostics_button = ttk.Button(
            controls,
            text="Save Error Report",
            command=self.save_diagnostics_report,
        )
        self.diagnostics_button.grid(row=0, column=13, sticky="e")

        output_frame = ttk.Frame(self.root, padding=(12, 0, 12, 8))
        output_frame.grid(row=1, column=0, sticky="ew")
        output_frame.columnconfigure(1, weight=1)
        output_frame.columnconfigure(5, weight=1)

        ttk.Label(output_frame, text="Autosave TXT").grid(row=0, column=0, sticky="w")
        output_entry = ttk.Entry(output_frame, textvariable=self.output_var)
        output_entry.grid(row=0, column=1, padx=(8, 8), sticky="ew")
        ttk.Button(output_frame, text="Browse", command=self.choose_output_path).grid(
            row=0, column=2, sticky="e"
        )

        ttk.Label(output_frame, text="Search").grid(row=0, column=3, padx=(18, 0), sticky="w")
        search_entry = ttk.Entry(output_frame, textvariable=self.search_var)
        search_entry.grid(row=0, column=4, padx=(8, 8), sticky="ew")
        search_entry.bind("<Return>", lambda _event: self.find_next_match())
        ttk.Button(output_frame, text="Next", command=self.find_next_match).grid(
            row=0, column=5, padx=(0, 8), sticky="w"
        )
        ttk.Button(output_frame, text="Clear", command=self.clear_search).grid(row=0, column=6, sticky="w")

        ttk.Label(output_frame, text="Log file").grid(row=1, column=0, sticky="w", pady=(8, 0))
        ttk.Entry(output_frame, textvariable=self.log_var, state="readonly").grid(
            row=1, column=1, padx=(8, 8), pady=(8, 0), sticky="ew"
        )
        ttk.Button(output_frame, text="Open Logs", command=self.open_logs_folder).grid(
            row=1, column=2, pady=(8, 0), sticky="e"
        )

        marker_frame = ttk.Frame(output_frame)
        marker_frame.grid(row=1, column=3, columnspan=4, pady=(8, 0), sticky="w")
        ttk.Label(marker_frame, text="Mark selected line").grid(row=0, column=0, sticky="w")
        ttk.Button(marker_frame, text=f"{IMPORTANT_ICON} Important", command=lambda: self.mark_selected_entry("IMPORTANT")).grid(
            row=0, column=1, padx=(8, 6), sticky="w"
        )
        ttk.Button(marker_frame, text=f"{TASK_ICON} Task", command=lambda: self.mark_selected_entry("TASK")).grid(
            row=0, column=2, padx=6, sticky="w"
        )
        ttk.Button(marker_frame, text=f"{QUESTION_ICON} Question", command=lambda: self.mark_selected_entry("QUESTION")).grid(
            row=0, column=3, padx=6, sticky="w"
        )
        ttk.Button(marker_frame, text="Clear Mark", command=lambda: self.mark_selected_entry(None)).grid(
            row=0, column=4, padx=6, sticky="w"
        )

        text_frame = ttk.Frame(self.root, padding=(12, 0, 12, 8))
        text_frame.grid(row=2, column=0, sticky="nsew")
        text_frame.rowconfigure(0, weight=1)
        text_frame.columnconfigure(0, weight=1)

        self.text_widget = tk.Text(
            text_frame,
            wrap="word",
            font=("Segoe UI", 11),
            undo=False,
        )
        self.text_widget.grid(row=0, column=0, sticky="nsew")
        self.text_widget.insert(
            "1.0",
            "Live transcript will appear here.\n"
            f"{PLAY_ICON} plays the line audio.\n"
            f"{TASK_ICON} marks a task, {QUESTION_ICON} marks a question, {IMPORTANT_ICON} marks an important line.\n"
            "Single-click a transcript line to select it, then click again or use Play Selected to hear that fragment.\n\n",
        )
        self.text_widget.bind("<ButtonRelease-1>", self.on_text_click)
        self.text_widget.tag_configure(SEARCH_TAG, background="#fff2a8")
        self.text_widget.tag_configure(CURRENT_SEARCH_TAG, background="#ffc94d")

        scrollbar = ttk.Scrollbar(text_frame, orient="vertical", command=self.text_widget.yview)
        scrollbar.grid(row=0, column=1, sticky="ns")
        self.text_widget.configure(yscrollcommand=scrollbar.set)

        status_bar = ttk.Frame(self.root, padding=(12, 0, 12, 12))
        status_bar.grid(row=3, column=0, sticky="ew")
        status_bar.columnconfigure(0, weight=1)

        ttk.Label(status_bar, textvariable=self.status_var).grid(row=0, column=0, sticky="w")

    def default_output_path(self) -> Path:
        return self.output_dir / f"meeting_{self.session_stamp}.txt"

    def default_audio_path(self, source_tag: str) -> Path:
        slug = slugify_source(source_tag)
        return self.audio_dir / f"meeting_{self.session_stamp}_{slug}.wav"

    def default_clip_path(self, source_tag: str, entry_id: int) -> Path:
        slug = slugify_source(source_tag)
        return self.audio_clips_dir / f"meeting_{self.session_stamp}_{slug}_{entry_id:05d}.wav"

    def choose_output_path(self) -> None:
        selected = filedialog.asksaveasfilename(
            title="Choose transcript TXT file",
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt")],
            initialfile=Path(self.output_var.get()).name,
            initialdir=str(Path(self.output_var.get()).parent),
        )
        if selected:
            self.output_var.set(selected)

    def start_transcription(self) -> None:
        if self.capture_thread and self.capture_thread.is_alive():
            return

        output_path = Path(self.output_var.get()).expanduser()
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text("", encoding="utf-8")

        self.session_stamp = time.strftime("%Y%m%d_%H%M%S")
        self.clear_runtime_state()
        self.stop_event.clear()
        self.text_widget.delete("1.0", tk.END)
        self.text_widget.insert("1.0", "")

        self.status_var.set("Loading model...")
        self.start_button.configure(state="disabled")
        self.stop_button.configure(state="normal")
        self.log_message(
            "info",
            (
                f"Starting transcription: language={self.language_var.get()} "
                f"model={self.model_var.get()} include_mic={self.include_mic_var.get()} "
                f"stable_mode={self.stable_mode_var.get()} "
                f"meeting_diarization={self.meeting_diarization_var.get()} "
                f"output={output_path}"
            ),
        )

        self.capture_thread = threading.Thread(target=self.capture_loop, daemon=True)
        self.worker_thread = threading.Thread(target=self.transcribe_loop, daemon=True)
        self.worker_thread.start()
        self.capture_thread.start()
        if self.include_mic_var.get():
            self.mic_thread = threading.Thread(target=self.microphone_loop, daemon=True)
            self.mic_thread.start()

    def stop_transcription(self) -> None:
        self.stop_event.set()
        self.start_button.configure(state="normal")
        self.stop_button.configure(state="disabled")
        self.status_var.set("Stopping...")
        self.log_message("info", "Stopping transcription")

    def clear_runtime_state(self) -> None:
        self.close_audio_wavefiles()
        self.last_emitted_by_source.clear()
        self.transcript_parts.clear()
        self.transcript_entries.clear()
        self.line_to_entry_id.clear()
        self.search_results.clear()
        self.current_search_index = -1
        self.selected_entry_id = None
        self.speaker_profiles.clear()
        self.next_speaker_id = 1
        self.audio_paths.clear()
        while not self.audio_queue.empty():
            try:
                self.audio_queue.get_nowait()
            except queue.Empty:
                break

    def get_record_seconds(self) -> int:
        return STABLE_RECORD_SECONDS if self.stable_mode_var.get() else DEFAULT_RECORD_SECONDS

    def get_block_frames(self) -> int:
        return STABLE_BLOCK_FRAMES if self.stable_mode_var.get() else DEFAULT_BLOCK_FRAMES

    def get_beam_size(self) -> int:
        return 1 if self.stable_mode_var.get() else 5

    def capture_loop(self) -> None:
        try:
            speaker = sc.default_speaker()
            if speaker is None:
                raise RuntimeError("No default speaker device found.")

            self.ui_queue.put(("status", f"Capturing audio from: {speaker.name}"))
            self.log_message("info", f"Default speaker: {speaker.name}")
            loopback_device = sc.get_microphone(
                id=str(speaker.id),
                include_loopback=True,
            )
            self.log_message("info", f"Loopback device: {loopback_device.name}")
            self.capture_source(loopback_device, channels=2, source_tag="Meeting")
        except Exception as exc:
            self.report_exception("Audio capture failed", exc)
            self.stop_event.set()

    def microphone_loop(self) -> None:
        try:
            microphone = sc.default_microphone()
            if microphone is None:
                raise RuntimeError("No default microphone device found.")

            self.ui_queue.put(("status", f"Microphone enabled: {microphone.name}"))
            self.log_message("info", f"Default microphone: {microphone.name}")
            self.capture_source(microphone, channels=1, source_tag="You")
        except Exception as exc:
            self.report_exception("Microphone capture failed", exc)
            self.stop_event.set()

    def ensure_audio_wavefile(self, source_tag: str) -> Path:
        with self.audio_write_lock:
            if source_tag not in self.audio_wavefiles:
                audio_path = self.default_audio_path(source_tag)
                wav_handle = wave.open(str(audio_path), "wb")
                wav_handle.setnchannels(1)
                wav_handle.setsampwidth(2)
                wav_handle.setframerate(SAMPLE_RATE)
                self.audio_wavefiles[source_tag] = wav_handle
                self.audio_paths[source_tag] = audio_path
                self.log_message("info", f"Audio recording enabled: {audio_path}")
            return self.audio_paths[source_tag]

    def append_audio_chunk(self, source_tag: str, chunk: np.ndarray) -> Path:
        audio_path = self.ensure_audio_wavefile(source_tag)
        with self.audio_write_lock:
            self.audio_wavefiles[source_tag].writeframes(pcm16_bytes(chunk))
        return audio_path

    def close_audio_wavefiles(self) -> None:
        with self.audio_write_lock:
            for wav_handle in self.audio_wavefiles.values():
                try:
                    wav_handle.close()
                except Exception:
                    pass
            self.audio_wavefiles.clear()

    def capture_source(self, device, channels: int, source_tag: str) -> None:
        source_elapsed = 0.0
        block_frames = self.get_block_frames()
        target_frames = SAMPLE_RATE * self.get_record_seconds()
        try:
            with device.recorder(
                samplerate=SAMPLE_RATE,
                channels=channels,
                blocksize=block_frames,
            ) as recorder:
                buffer: list[np.ndarray] = []
                frames_collected = 0

                while not self.stop_event.is_set():
                    data = recorder.record(numframes=block_frames)
                    if data.size == 0:
                        continue

                    if data.ndim == 1:
                        mono = data.astype(np.float32)
                    else:
                        mono = data.mean(axis=1).astype(np.float32)
                    buffer.append(mono)
                    frames_collected += len(mono)

                    if frames_collected >= target_frames:
                        chunk = np.concatenate(buffer)
                        buffer.clear()
                        frames_collected = 0
                        audio_start = source_elapsed
                        audio_end = audio_start + (len(chunk) / SAMPLE_RATE)
                        source_elapsed = audio_end
                        audio_path = self.append_audio_chunk(source_tag, chunk)
                        self.enqueue_audio(source_tag, chunk, audio_start, audio_end, audio_path)

                if buffer:
                    chunk = np.concatenate(buffer)
                    audio_start = source_elapsed
                    audio_end = audio_start + (len(chunk) / SAMPLE_RATE)
                    self.append_audio_chunk(source_tag, chunk)
                    self.enqueue_audio(
                        source_tag,
                        chunk,
                        audio_start,
                        audio_end,
                        self.audio_paths.get(source_tag),
                    )
        finally:
            with self.audio_write_lock:
                wav_handle = self.audio_wavefiles.pop(source_tag, None)
                if wav_handle is not None:
                    try:
                        wav_handle.close()
                    except Exception:
                        pass

    def enqueue_audio(
        self,
        source_tag: str,
        chunk: np.ndarray,
        audio_start: float,
        audio_end: float,
        audio_path: Path | None,
    ) -> None:
        if self.stop_event.is_set():
            return

        peak = float(np.max(np.abs(chunk))) if chunk.size else 0.0
        if peak < 0.008:
            return

        try:
            self.audio_queue.put((source_tag, chunk, audio_start, audio_end, audio_path), timeout=1)
        except queue.Full:
            self.ui_queue.put(("status", "Transcriber is busy. Dropping an audio chunk."))

    def transcribe_loop(self) -> None:
        try:
            language = None if self.language_var.get() == "auto" else self.language_var.get()
            model_name = self.model_var.get()

            self.model = WhisperModel(model_name, device="cpu", compute_type="int8")
            self.load_diarization_encoder_if_needed()
            self.ui_queue.put(("status", f"Model loaded: {model_name}. Listening..."))
            self.log_message("info", f"Whisper model loaded: {model_name}")

            while not self.stop_event.is_set() or not self.audio_queue.empty():
                try:
                    source_tag, chunk, audio_start, audio_end, audio_path = self.audio_queue.get(timeout=0.5)
                except queue.Empty:
                    continue

                segments, _ = self.model.transcribe(
                    chunk,
                    language=language,
                    vad_filter=False,
                    beam_size=self.get_beam_size(),
                    condition_on_previous_text=False,
                    word_timestamps=True,
                )
                segment_list = list(segments)
                if not segment_list:
                    continue

                if self.should_diarize_source(source_tag):
                    self.add_diarized_entries(segment_list, chunk, source_tag)
                    continue

                self.add_segment_entries(segment_list, chunk, source_tag)

            self.ui_queue.put(("status", "Stopped"))
            self.log_message("info", "Transcription loop stopped")
        except Exception as exc:
            self.report_exception("Transcription failed", exc)
            self.stop_event.set()
        finally:
            self.close_audio_wavefiles()

    def load_diarization_encoder_if_needed(self) -> None:
        self.diarization_encoder = None
        if not self.meeting_diarization_var.get():
            return
        if VoiceEncoder is None or preprocess_wav is None:
            self.log_message("warning", "Meeting diarization requested, but resemblyzer is not installed.")
            self.ui_queue.put(
                ("status", "Speaker split is off: install resemblyzer to enable meeting diarization.")
            )
            return
        self.ui_queue.put(("status", "Loading diarization encoder..."))
        self.diarization_encoder = VoiceEncoder()
        self.log_message("info", "Diarization encoder loaded")

    def should_diarize_source(self, source_tag: str) -> bool:
        return (
            source_tag == "Meeting"
            and self.meeting_diarization_var.get()
            and self.diarization_encoder is not None
        )

    def add_diarized_entries(
        self,
        segments,
        chunk: np.ndarray,
        source_tag: str,
    ) -> None:
        added_lines = 0
        for segment in segments:
            text = normalize_text(segment.text)
            if not text:
                continue

            segment_start, segment_end = self.get_segment_audio_bounds(segment)
            speaker_tag = self.resolve_speaker_tag(chunk, segment_start, segment_end) or source_tag
            clip_path, clip_duration = self.create_entry_audio_clip(chunk, speaker_tag, segment_start, segment_end)
            self.add_transcript_entry(text, speaker_tag, 0.0, clip_duration, clip_path)
            added_lines += 1

        if added_lines:
            self.last_emitted_by_source[source_tag] = self.transcript_entries[-1].transcript_text

    def add_segment_entries(
        self,
        segments,
        chunk: np.ndarray,
        source_tag: str,
    ) -> None:
        for segment in segments:
            text = normalize_text(segment.text)
            if not text:
                continue
            if text == self.last_emitted_by_source.get(source_tag):
                continue

            segment_start, segment_end = self.get_segment_audio_bounds(segment)
            clip_path, clip_duration = self.create_entry_audio_clip(chunk, source_tag, segment_start, segment_end)
            self.last_emitted_by_source[source_tag] = text
            self.add_transcript_entry(text, source_tag, 0.0, clip_duration, clip_path)

    def get_segment_audio_bounds(self, segment) -> tuple[float, float]:
        segment_start = max(0.0, float(segment.start))
        segment_end = max(segment_start, float(segment.end))
        words = getattr(segment, "words", None)
        if not words:
            return segment_start, segment_end

        timed_words = [
            word for word in words
            if getattr(word, "start", None) is not None and getattr(word, "end", None) is not None
        ]
        if not timed_words:
            return segment_start, segment_end

        word_start = max(0.0, float(timed_words[0].start))
        word_end = max(word_start, float(timed_words[-1].end))
        return word_start, word_end

    def create_entry_audio_clip(
        self,
        chunk: np.ndarray,
        source_tag: str,
        segment_start: float,
        segment_end: float,
    ) -> tuple[Path, float]:
        padded_start = max(0.0, segment_start - CLIP_PADDING_BEFORE)
        padded_end = min(len(chunk) / SAMPLE_RATE, segment_end + CLIP_PADDING_AFTER)
        start_idx = max(0, int(padded_start * SAMPLE_RATE))
        end_idx = min(len(chunk), max(start_idx + 1, int(padded_end * SAMPLE_RATE)))
        clip_chunk = chunk[start_idx:end_idx]
        if clip_chunk.size == 0:
            clip_chunk = chunk[max(0, start_idx - 1) : max(start_idx, start_idx + 1)]
        entry_id = len(self.transcript_entries)
        clip_path = self.default_clip_path(source_tag, entry_id)
        write_wav_file(clip_path, clip_chunk)
        return clip_path, len(clip_chunk) / SAMPLE_RATE

    def resolve_speaker_tag(self, chunk: np.ndarray, segment_start: float, segment_end: float) -> str | None:
        if self.diarization_encoder is None or preprocess_wav is None:
            return None

        duration = segment_end - segment_start
        if duration < MIN_DIARIZATION_SECONDS:
            return None

        start_idx = max(0, int(segment_start * SAMPLE_RATE))
        end_idx = min(len(chunk), max(start_idx + 1, int(segment_end * SAMPLE_RATE)))
        segment_audio = chunk[start_idx:end_idx]
        if len(segment_audio) < int(MIN_DIARIZATION_SECONDS * SAMPLE_RATE):
            return None

        try:
            processed = preprocess_wav(segment_audio.astype(np.float32))
            embedding = self.diarization_encoder.embed_utterance(processed)
        except Exception as exc:
            self.log_message("warning", f"Speaker embedding failed: {exc}")
            return None

        return self.assign_speaker_label(embedding)

    def assign_speaker_label(self, embedding: np.ndarray) -> str:
        best_label: str | None = None
        best_score = -1.0

        for label, (centroid, sample_count) in self.speaker_profiles.items():
            denominator = np.linalg.norm(centroid) * np.linalg.norm(embedding)
            if denominator <= 0:
                continue
            score = float(np.dot(centroid, embedding) / denominator)
            if score > best_score:
                best_score = score
                best_label = label

        if best_label is None or best_score < SPEAKER_MATCH_THRESHOLD:
            best_label = f"Speaker {self.next_speaker_id}"
            self.next_speaker_id += 1
            self.speaker_profiles[best_label] = (embedding, 1)
            return best_label

        centroid, sample_count = self.speaker_profiles[best_label]
        updated = ((centroid * sample_count) + embedding) / (sample_count + 1)
        self.speaker_profiles[best_label] = (updated, sample_count + 1)
        return best_label

    def add_transcript_entry(
        self,
        text: str,
        source_tag: str,
        audio_start: float,
        audio_end: float,
        audio_path: Path | None,
    ) -> None:
        entry = TranscriptEntry(
            entry_id=len(self.transcript_entries),
            source_tag=source_tag,
            transcript_text=text,
            created_at=timestamp_now(),
            audio_path=audio_path,
            audio_start=audio_start,
            audio_end=audio_end,
        )
        self.transcript_entries.append(entry)
        display = entry.display_text(self.timestamps_var.get())
        self.transcript_parts.append(display)
        self.append_to_output_file(display)
        self.ui_queue.put(("text", str(entry.entry_id)))

    def append_to_output_file(self, text: str) -> None:
        output_path = Path(self.output_var.get()).expanduser()
        with output_path.open("a", encoding="utf-8") as handle:
            handle.write(text + "\n")

    def render_entry(self, entry: TranscriptEntry) -> None:
        display = entry.display_text(self.timestamps_var.get())
        line_number = entry.line_number
        self.transcript_parts[entry.entry_id] = display
        if line_number is None:
            line_number = int(self.text_widget.index("end-1c").split(".")[0])
            if self.text_widget.get("1.0", "end-1c"):
                line_number += 1
            entry.line_number = line_number
            self.line_to_entry_id[line_number] = entry.entry_id
            self.text_widget.insert(tk.END, display + "\n")
        else:
            self.text_widget.delete(f"{line_number}.0", f"{line_number}.end")
            self.text_widget.insert(f"{line_number}.0", display)

        self.apply_line_styles(entry)
        self.apply_line_action_tags(entry)
        self.text_widget.see(f"{entry.line_number}.0")

    def apply_line_styles(self, entry: TranscriptEntry) -> None:
        if entry.line_number is None:
            return

        tag_name = f"{LINE_PREFIX}{entry.entry_id}"
        line_start = f"{entry.line_number}.0"
        line_end = f"{entry.line_number}.end"
        self.text_widget.tag_remove(tag_name, "1.0", tk.END)

        background = ""
        foreground = ""
        if entry.marker == "IMPORTANT":
            background = "#ffe7a8"
        elif entry.marker == "TASK":
            background = "#dff5dc"
        elif entry.marker == "QUESTION":
            background = "#dcecff"

        if self.selected_entry_id == entry.entry_id:
            foreground = "#0b4f8a"

        self.text_widget.tag_add(tag_name, line_start, line_end)
        self.text_widget.tag_configure(tag_name, background=background, foreground=foreground)

    def apply_line_action_tags(self, entry: TranscriptEntry) -> None:
        if entry.line_number is None:
            return

        line = entry.line_number
        prefix = action_prefix()
        positions = {
            "play": (0, len(PLAY_ICON)),
            "task": (len(PLAY_ICON) + 1, len(PLAY_ICON) + 1 + len(TASK_ICON)),
            "question": (
                len(PLAY_ICON) + 1 + len(TASK_ICON) + 1,
                len(PLAY_ICON) + 1 + len(TASK_ICON) + 1 + len(QUESTION_ICON),
            ),
            "important": (
                len(PLAY_ICON) + 1 + len(TASK_ICON) + 1 + len(QUESTION_ICON) + 1,
                len(prefix),
            ),
        }

        for action_name, (start_col, end_col) in positions.items():
            tag_name = f"action_{action_name}_{entry.entry_id}"
            start_idx = f"{line}.{start_col}"
            end_idx = f"{line}.{end_col}"
            self.text_widget.tag_remove(tag_name, "1.0", tk.END)
            self.text_widget.tag_add(tag_name, start_idx, end_idx)
            self.text_widget.tag_bind(
                tag_name,
                "<Button-1>",
                lambda _event, entry_id=entry.entry_id, action=action_name: self.handle_inline_action(
                    entry_id, action
                ),
            )
            if action_name == "play":
                self.text_widget.tag_configure(tag_name, foreground="#0b63c9", underline=True)
            elif action_name == "task":
                color = "#1f7a1f" if entry.marker == "TASK" else "#666666"
                self.text_widget.tag_configure(tag_name, foreground=color)
            elif action_name == "question":
                color = "#8a5a00" if entry.marker == "QUESTION" else "#666666"
                self.text_widget.tag_configure(tag_name, foreground=color)
            else:
                color = "#9b1c1c" if entry.marker == "IMPORTANT" else "#666666"
                self.text_widget.tag_configure(tag_name, foreground=color)

    def handle_inline_action(self, entry_id: int, action_name: str) -> str:
        entry = self.transcript_entries[entry_id]
        self.selected_entry_id = entry_id
        if action_name == "play":
            self.play_entry(entry)
        elif action_name == "task":
            entry.marker = None if entry.marker == "TASK" else "TASK"
            self.render_entry(entry)
            self.apply_search_highlights()
            self.status_var.set("Task mark updated.")
        elif action_name == "question":
            entry.marker = None if entry.marker == "QUESTION" else "QUESTION"
            self.render_entry(entry)
            self.apply_search_highlights()
            self.status_var.set("Question mark updated.")
        elif action_name == "important":
            entry.marker = None if entry.marker == "IMPORTANT" else "IMPORTANT"
            self.render_entry(entry)
            self.apply_search_highlights()
            self.status_var.set("Important mark updated.")
        return "break"

    def refresh_transcript_view(self) -> None:
        self.text_widget.delete("1.0", tk.END)
        self.line_to_entry_id.clear()
        for entry in self.transcript_entries:
            entry.line_number = None
            self.render_entry(entry)
        self.apply_search_highlights()

    def on_text_click(self, event: tk.Event) -> None:
        index = self.text_widget.index(f"@{event.x},{event.y}")
        current_tags = self.text_widget.tag_names(index)
        if any(tag.startswith("action_") for tag in current_tags):
            return
        line_number = int(index.split(".")[0])
        entry_id = self.line_to_entry_id.get(line_number)
        if entry_id is None:
            return

        if self.selected_entry_id == entry_id:
            self.play_entry(self.transcript_entries[entry_id])
            return

        self.selected_entry_id = entry_id
        for entry in self.transcript_entries:
            if entry.line_number is not None:
                self.apply_line_styles(entry)
        self.status_var.set(
            "Line selected. Click the same line again or use Play Selected to hear the fragment."
        )

    def play_selected_entry(self) -> None:
        if self.selected_entry_id is None:
            messagebox.showinfo(APP_TITLE, "Select a transcript line first.")
            return
        self.play_entry(self.transcript_entries[self.selected_entry_id])

    def play_entry(self, entry: TranscriptEntry) -> None:
        if not entry.has_audio():
            messagebox.showinfo(APP_TITLE, "Audio for this line is not available yet.")
            return

        if self.playback_thread and self.playback_thread.is_alive():
            self.status_var.set("Audio playback is already running.")
            return

        self.playback_thread = threading.Thread(target=self.play_audio_fragment, args=(entry,), daemon=True)
        self.playback_thread.start()

    def play_audio_fragment(self, entry: TranscriptEntry) -> None:
        self.ui_queue.put(("status", f"Playing audio for line {entry.entry_id + 1}..."))
        temp_path: Path | None = None
        try:
            start_frame = max(0, int(entry.audio_start * SAMPLE_RATE))
            end_frame = max(start_frame + 1, int(entry.audio_end * SAMPLE_RATE))

            with wave.open(str(entry.audio_path), "rb") as source_handle:
                total_frames = source_handle.getnframes()
                start_frame = min(start_frame, total_frames)
                end_frame = min(max(end_frame, start_frame + 1), total_frames)
                source_handle.setpos(start_frame)
                audio_bytes = source_handle.readframes(end_frame - start_frame)

            if not audio_bytes:
                self.ui_queue.put(("status", "No audio data found for the selected line."))
                return

            with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as temp_file:
                temp_path = Path(temp_file.name)

            with wave.open(str(temp_path), "wb") as fragment_handle:
                fragment_handle.setnchannels(1)
                fragment_handle.setsampwidth(2)
                fragment_handle.setframerate(SAMPLE_RATE)
                fragment_handle.writeframes(audio_bytes)

            if platform.system() == "Windows":
                self.play_audio_fragment_windows(audio_bytes, temp_path)
            elif platform.system() == "Darwin":
                subprocess.run(["afplay", str(temp_path)], check=False)
            else:
                subprocess.run(["aplay", str(temp_path)], check=False)

            self.ui_queue.put(("status", "Playback finished."))
        except Exception as exc:
            self.report_exception("Audio playback failed", exc)
        finally:
            if temp_path is not None:
                try:
                    temp_path.unlink(missing_ok=True)
                except Exception:
                    pass

    def play_audio_fragment_windows(self, audio_bytes: bytes, temp_path: Path) -> None:
        if sa is not None:
            try:
                play_object = sa.play_buffer(audio_bytes, 1, 2, SAMPLE_RATE)
                play_object.wait_done()
                return
            except Exception as exc:
                self.log_message("warning", f"simpleaudio playback failed for {temp_path}: {exc}")

        import winsound

        try:
            winsound.PlaySound(str(temp_path), winsound.SND_FILENAME)
            return
        except RuntimeError as exc:
            self.log_message("warning", f"winsound playback failed for {temp_path}: {exc}")

        ffplay_result = subprocess.run(
            ["ffplay", "-nodisp", "-autoexit", "-loglevel", "error", str(temp_path)],
            capture_output=True,
            text=True,
            check=False,
        )
        if ffplay_result.returncode == 0:
            return
        stderr = ffplay_result.stderr.strip() or ffplay_result.stdout.strip()
        if stderr:
            self.log_message("warning", f"ffplay playback failed for {temp_path}: {stderr}")

        escaped_temp_path = str(temp_path).replace("'", "''")
        powershell_command = (
            "Add-Type -AssemblyName presentationCore; "
            "$player = New-Object System.Media.SoundPlayer "
            f"'{escaped_temp_path}'; "
            "$player.Load(); "
            "$player.PlaySync();"
        )
        result = subprocess.run(
            ["powershell", "-NoProfile", "-Command", powershell_command],
            capture_output=True,
            text=True,
            check=False,
        )
        if result.returncode != 0:
            stderr = result.stderr.strip() or result.stdout.strip() or "Unknown playback error"
            raise RuntimeError(stderr)

    def mark_selected_entry(self, marker: str | None) -> None:
        if self.selected_entry_id is None:
            messagebox.showinfo(APP_TITLE, "Select a transcript line first.")
            return

        entry = self.transcript_entries[self.selected_entry_id]
        entry.marker = marker
        self.render_entry(entry)
        self.apply_search_highlights()
        self.status_var.set("Line updated.")

    def clear_search(self) -> None:
        self.search_var.set("")
        self.search_results.clear()
        self.current_search_index = -1
        self.text_widget.tag_remove(SEARCH_TAG, "1.0", tk.END)
        self.text_widget.tag_remove(CURRENT_SEARCH_TAG, "1.0", tk.END)
        self.status_var.set("Search cleared.")

    def apply_search_highlights(self) -> None:
        query = self.search_var.get().strip()
        self.text_widget.tag_remove(SEARCH_TAG, "1.0", tk.END)
        self.text_widget.tag_remove(CURRENT_SEARCH_TAG, "1.0", tk.END)
        self.search_results.clear()
        self.current_search_index = -1

        if not query:
            return

        start = "1.0"
        while True:
            match_start = self.text_widget.search(query, start, stopindex=tk.END, nocase=True)
            if not match_start:
                break
            match_end = f"{match_start}+{len(query)}c"
            self.text_widget.tag_add(SEARCH_TAG, match_start, match_end)
            self.search_results.append(match_start)
            start = match_end

        if self.search_results:
            self.current_search_index = 0
            current = self.search_results[0]
            current_end = f"{current}+{len(query)}c"
            self.text_widget.tag_add(CURRENT_SEARCH_TAG, current, current_end)
            self.text_widget.see(current)
            self.status_var.set(f"Found {len(self.search_results)} matches.")
        else:
            self.status_var.set("No matches found.")

    def find_next_match(self) -> None:
        query = self.search_var.get().strip()
        if not query:
            self.status_var.set("Enter text to search.")
            return

        if not self.search_results:
            self.apply_search_highlights()
            return

        self.text_widget.tag_remove(CURRENT_SEARCH_TAG, "1.0", tk.END)
        self.current_search_index = (self.current_search_index + 1) % len(self.search_results)
        current = self.search_results[self.current_search_index]
        current_end = f"{current}+{len(query)}c"
        self.text_widget.tag_add(CURRENT_SEARCH_TAG, current, current_end)
        self.text_widget.see(current)
        self.status_var.set(
            f"Match {self.current_search_index + 1} of {len(self.search_results)}."
        )

    def build_summary_text(self) -> str:
        if not self.transcript_entries:
            return "No transcript yet."

        important_lines = [entry for entry in self.transcript_entries if entry.marker == "IMPORTANT"]
        task_lines = [entry for entry in self.transcript_entries if entry.marker == "TASK"]
        question_lines = [entry for entry in self.transcript_entries if entry.marker == "QUESTION"]

        if not important_lines:
            important_lines = self.transcript_entries[: min(5, len(self.transcript_entries))]

        inferred_questions = [
            entry for entry in self.transcript_entries if "?" in entry.transcript_text and entry not in question_lines
        ][:5]

        lines = [
            "Meeting Summary",
            f"Generated: {timestamp_now()}",
            f"Transcript lines: {len(self.transcript_entries)}",
            "",
            "Key points:",
        ]
        lines.extend(
            f"- {entry.transcript_text}" for entry in important_lines[:7]
        )

        lines.append("")
        lines.append("Tasks:")
        if task_lines:
            lines.extend(f"- {entry.transcript_text}" for entry in task_lines[:10])
        else:
            lines.append("- No lines marked as tasks yet.")

        lines.append("")
        lines.append("Questions:")
        if question_lines or inferred_questions:
            for entry in (question_lines + inferred_questions)[:10]:
                lines.append(f"- {entry.transcript_text}")
        else:
            lines.append("- No questions detected.")

        lines.append("")
        lines.append("Audio files:")
        if self.audio_paths:
            for source_tag, audio_path in self.audio_paths.items():
                lines.append(f"- {source_tag}: {audio_path}")
        else:
            lines.append("- No audio files saved.")

        return "\n".join(lines)

    def show_summary(self) -> None:
        summary_text = self.build_summary_text()
        summary_window = tk.Toplevel(self.root)
        summary_window.title("Meeting Summary")
        summary_window.geometry("760x560")

        frame = ttk.Frame(summary_window, padding=12)
        frame.pack(fill="both", expand=True)

        summary_widget = tk.Text(frame, wrap="word", font=("Segoe UI", 10))
        summary_widget.pack(fill="both", expand=True)
        summary_widget.insert("1.0", summary_text)

        actions = ttk.Frame(frame)
        actions.pack(fill="x", pady=(10, 0))

        def save_summary() -> None:
            suggested = Path(self.output_var.get()).with_name(
                f"{Path(self.output_var.get()).stem}_summary.txt"
            )
            selected = filedialog.asksaveasfilename(
                title="Save summary",
                defaultextension=".txt",
                filetypes=[("Text files", "*.txt")],
                initialfile=suggested.name,
                initialdir=str(suggested.parent),
            )
            if not selected:
                return
            Path(selected).write_text(summary_widget.get("1.0", "end-1c") + "\n", encoding="utf-8")
            self.status_var.set(f"Summary saved: {selected}")

        ttk.Button(actions, text="Save Summary", command=save_summary).pack(side="left")
        ttk.Button(actions, text="Close", command=summary_window.destroy).pack(side="right")

    def export_docx(self) -> None:
        if not self.transcript_entries:
            messagebox.showinfo(APP_TITLE, "There is no transcript yet.")
            return

        suggested = Path(self.output_var.get()).with_suffix(".docx")
        selected = filedialog.asksaveasfilename(
            title="Export DOCX",
            defaultextension=".docx",
            filetypes=[("Word document", "*.docx")],
            initialfile=suggested.name,
            initialdir=str(suggested.parent),
        )
        if not selected:
            return

        document = Document()
        document.add_heading("Live Meeting Transcript", level=1)
        document.add_paragraph(f"Created: {timestamp_now()}")
        for entry in self.transcript_entries:
            document.add_paragraph(entry.display_text(self.timestamps_var.get()))
        summary_text = self.build_summary_text()
        document.add_page_break()
        document.add_heading("Meeting Summary", level=1)
        for line in summary_text.splitlines():
            document.add_paragraph(line)
        document.save(selected)
        self.log_message("info", f"DOCX exported: {selected}")
        messagebox.showinfo(APP_TITLE, f"Saved DOCX:\n{selected}")

    def open_logs_folder(self) -> None:
        logs_dir = self.logs_dir.resolve()
        try:
            if platform.system() == "Windows":
                import os

                os.startfile(logs_dir)  # type: ignore[attr-defined]
            else:
                messagebox.showinfo(APP_TITLE, f"Logs folder:\n{logs_dir}")
        except Exception as exc:
            self.report_exception("Failed to open logs folder", exc)

    def log_message(self, level: str, message: str) -> None:
        line = f"[{timestamp_now()}] [{level.upper()}] {message}\n"
        with self.log_path.open("a", encoding="utf-8") as handle:
            handle.write(line)

    def report_exception(self, context: str, exc: Exception) -> None:
        details = f"{context}: {exc}\n{traceback.format_exc()}"
        self.last_error_details = details
        self.log_message("error", details)
        self.ui_queue.put(("error", f"{context}: {exc}"))

    def build_diagnostics_text(self) -> str:
        lines = [
            f"Generated: {timestamp_now()}",
            f"Python: {platform.python_version()}",
            f"Platform: {platform.platform()}",
            f"System: {platform.system()} {platform.release()}",
            f"App: {APP_TITLE}",
            f"Model: {self.model_var.get()}",
            f"Language: {self.language_var.get()}",
            f"Include microphone: {self.include_mic_var.get()}",
            f"Output file: {self.output_var.get()}",
            f"Log file: {self.log_path.resolve()}",
            "Audio files:",
        ]
        if self.audio_paths:
            lines.extend(f"- {source_tag}: {path}" for source_tag, path in self.audio_paths.items())
        else:
            lines.append("- No audio saved yet")
        lines.extend(
            [
                "",
                "Last error:",
                self.last_error_details or "No captured error",
                "",
                "Recent transcript preview:",
                "\n".join(
                    entry.display_text(self.timestamps_var.get()) for entry in self.transcript_entries[-20:]
                )
                or "No transcript yet",
            ]
        )
        return "\n".join(lines) + "\n"

    def save_diagnostics_report(self) -> None:
        suggested = self.logs_dir / f"error_report_{time.strftime('%Y%m%d_%H%M%S')}.txt"
        selected = filedialog.asksaveasfilename(
            title="Save error report",
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt")],
            initialfile=suggested.name,
            initialdir=str(suggested.parent),
        )
        if not selected:
            return

        report_path = Path(selected)
        report_path.write_text(self.build_diagnostics_text(), encoding="utf-8")
        self.log_message("info", f"Diagnostics report saved: {report_path}")
        messagebox.showinfo(
            APP_TITLE,
            "Report saved.\nSend me this file together with the latest log file if something breaks.",
        )

    def process_ui_queue(self) -> None:
        while True:
            try:
                event_type, payload = self.ui_queue.get_nowait()
            except queue.Empty:
                break

            if event_type == "text":
                entry = self.transcript_entries[int(payload)]
                self.render_entry(entry)
                self.apply_search_highlights()
                self.status_var.set("Listening...")
            elif event_type == "status":
                self.status_var.set(payload)
            elif event_type == "error":
                self.status_var.set(payload)
                self.start_button.configure(state="normal")
                self.stop_button.configure(state="disabled")
                messagebox.showerror(
                    APP_TITLE,
                    f"{payload}\n\nLog file:\n{self.log_path}\n\n"
                    "Use 'Save Error Report' and send me the report plus the log file.",
                )

        self.root.after(150, self.process_ui_queue)

    def on_close(self) -> None:
        self.stop_event.set()
        self.close_audio_wavefiles()
        self.log_message("info", "Application closing")
        self.root.after(100, self.root.destroy)


def main() -> None:
    root = tk.Tk()
    style = ttk.Style()
    if "clam" in style.theme_names():
        style.theme_use("clam")
    LiveTranscriberApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
